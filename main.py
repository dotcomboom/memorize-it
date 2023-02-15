import os
import re
from io import BytesIO

from docx import Document # python-docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from flask import Flask, redirect, render_template, request, send_file, url_for


def rewrite(filename, mode):    
    doc = Document(filename)
    # Loop through the paragraphs.
    for para in doc.paragraphs:
        # Loop through the runs.
        for run in para.runs:
            # If the run is bolded or italicized, don't change it.
            if run.bold or run.italic:
                continue
            # Check if mode is "fill" or "strip".
            if mode == 'fill':
                # Replace each word in the run with the first letter of the word, and then the character _ for each letter after.
                run.text = re.sub(r'\w+', lambda m: m.group(0)[0] + '_' * (len(m.group(0)) - 1), run.text)
            elif mode == 'strip':
                # Replace each word in the run with the first letter of the word.
                run.text = re.sub(r'\w+', lambda m: m.group(0)[0], run.text)
    # Return the document.
    return doc

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

# post /upload to run rewrite_from_file
@app.route('/upload', methods=['POST'])
def upload():
    if request.method != 'POST':
        return
    file = request.files['file']
    if file:
        # run rewrite_from_file
        # run with both modes, then combine the documents
        ogdoc = rewrite(file, '')
        filled = rewrite(file, 'fill')
        stripped = rewrite(file, 'strip')
        # add the contents of doc1 and doc2 to a new document
        memorizable = Document()

        # set up Monospace style
        styles = memorizable.styles
        style = styles.add_style('Monospace', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Courier New'
        style.font.size = Pt(16)

        def copy_doc(source,dest):
            for para in source.paragraphs:
                np = dest.add_paragraph()
                for run in para.runs:
                    nr = np.add_run(run.text)
                    if run.bold:
                        nr.bold = True
                    if run.italic:
                        nr.italic = True
                    if run.underline:
                        nr.underline = True
                # if request.form.get('mono') == 'on':
                np.style = 'Monospace'

        sections = 0
        #if request.form.get('instructions') == 'on': sections += 1
        if request.form.get('first_letter') == 'on': sections += 1
        if request.form.get('fillable') == 'on': sections += 1
        if request.form.get('original') == 'on': sections += 1

        #if request.form.get('instructions') == 'on':
        #    memorizable.add_heading('Directions', 0)
        #    memorizable.add_paragraph('This packet created for the text: ').add_run(file.filename).bold = True
        #    memorizable.add_paragraph('This packet consists of three versions of your document.')
            # add numbered bullet list in bold
        #    memorizable.add_paragraph('A version of the text with each word reduced to the first letter. Your goal is to be able to read entirely off of this after a bit of practice, and then, completely from memory!', style='List Bullet')
        #    memorizable.add_paragraph('A version of the text with each word in a fill in the blank format with the first letter shown. This is so you can see the length of each word (including punctuation), or easily peek at the original when you\'re practicing. You don\'t need to fill it in, but you can if you want to quiz yourself!', style='List Bullet')
        #    memorizable.add_paragraph('The original text.', style='List Bullet')
        #    memorizable.add_paragraph('It might seem overwhelming at first! Just allow yourself to peek at the original text while you\'re working on memorization. Good luck!')

            # memorizable.add_paragraph('Feel free to skip printing these introductory instructions or printing the page ranges you need. Each version of the text is separated by page breaks.')

        #    if request.form.get('mono') == 'on':
        #        memorizable.add_paragraph('The Monospace option has been chosen, so the text is in monospace, 16pt font. Feel free to change it with Word as you wish.')
          #      else:
           #         memorizable.add_paragraph('The text has kept its original font as close as possible. If desired, you can change the font to monospace and/or increase the size for readability, or reupload your document with the Monospace option turned on.')
        memorizable.sections[0].footer.add_paragraph('{}'.format(request.host))

           # sections -= 1
         #   if sections > 0:
         #       memorizable.add_page_break()

        #if request.form.get('first_letter') == 'on':
        memorizable.add_heading('Original Text', 0)
        copy_doc(ogdoc, memorizable)

        sections -= 1
        if sections > 0:
            memorizable.add_page_break()

        memorizable.add_heading('First-Letter', 0)
        copy_doc(stripped, memorizable)

        sections -= 1
        if sections > 0:
            memorizable.add_page_break()   

        memorizable.add_heading('Fillable', 0)
        copy_doc(filled, memorizable)

        sections -= 1
        if sections > 0:
            memorizable.add_page_break() 

        # save the new document

        # mono = '_mono' if request.form.get('mono') == 'on' else ''
        mono = ''

        # get memorizable docx to be sent to the user without being saved to disk


        filename = './.temp/{}_memorizable{}.docx'.format(file.filename.replace('.docx', ''), mono)

        # if temp folder doesn't exist, create it
        if not os.path.exists('./.temp'):
            os.makedirs('./.temp')

        # if file already exists, delete it
        if os.path.exists(filename):
            os.remove(filename)

        memorizable.save(filename)

        # read the bytes of the file into memory
        with open(filename, 'rb') as f:
            memorizable_bytes = f.read()
            f.close()
        os.remove(filename)
        # return the file
        return send_file(BytesIO(memorizable_bytes), as_attachment=True, download_name=file.filename.replace('.docx', '_memorizable{}.docx'.format(mono)))

    else:
        return redirect(url_for('index'))
            
# run the app
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080)
